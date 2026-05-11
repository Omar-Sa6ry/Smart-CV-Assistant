import sys
import io
import os
import joblib
import numpy as np
import pdfplumber
from docx import Document
import warnings

# Fix encoding
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

warnings.filterwarnings('ignore')

# ── Import Engine ─────────────────────────────────────────────────────────────
try:
    from engine import analyze_cv_engine, SKILLS
    from train_model_v2 import normalize_text, NORMALIZE_MAP
except ImportError:
    # If run from outside directory
    sys.path.append(os.path.dirname(__file__))
    from engine import analyze_cv_engine, SKILLS
    from train_model_v2 import normalize_text, NORMALIZE_MAP

# ── Paths ─────────────────────────────────────────────────────────────────────
MODEL_PATH   = 'resume_model_v2.pkl'
ENCODER_PATH = 'label_encoder_v2.pkl'
CV_PATH      = 'Omar Ahmed Sabry (Backend Developer).pdf'

def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    try:
        if ext == ".pdf":
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Use layout=True to better preserve structure in multi-column CVs
                    t = page.extract_text(layout=True)
                    if t: text += t + "\n"
        elif ext == ".docx":
            doc = Document(file_path)
            # Extract from paragraphs
            text_parts = [p.text for p in doc.paragraphs]
            # Extract from tables (very common in CVs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            text = "\n".join(text_parts)
        elif ext == ".txt":
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            # For .doc or other formats, we might need other libraries, 
            # but for now we report it's not supported or try reading as text
            return None
            
        return text if text.strip() else None
    except Exception as e:
        print(f"FAILED EXTRACTION: {e}"); return None

def run_cv_audit(file_path):
    if not os.path.exists(file_path):
        print(f"FAILED: File not found: {file_path}"); return

    print(f"\nAnalyzing: {file_path}")
    text = extract_text(file_path)
    if not text: return

    if not all(os.path.exists(p) for p in [MODEL_PATH, ENCODER_PATH]):
        print("ERROR: Model files missing."); return

    try:
        model = joblib.load(MODEL_PATH)
        le    = joblib.load(ENCODER_PATH)

        # Predict
        norm_text = normalize_text(text)
        probs = model.predict_proba([norm_text])[0]
        pred_id = np.argmax(probs)
        role = le.inverse_transform([pred_id])[0]
        conf = probs[pred_id] * 100

        # Analysis
        res = analyze_cv_engine(text, role, normalize_text, NORMALIZE_MAP)

        W = "=" * 68
        print(f"\n{W}")
        print(f" ATS RESUME AUDIT REPORT (v2-Ultra)")
        print(W)
        print(f"  Predicted Role   : {role} ({conf:.1f}%)")
        print(f"  Experience       : {res['exp_years']} years | {res['seniority']}")
        print(f"  Overall ATS Score: {res['overallScore']}/100")
        print(W)

        print(f"\n🔑 KEYWORD MATCH: {res['kw_score']}%")
        print(f"  ✅ Found: {', '.join(res['found_skills'][:15])}...")
        print(f"  🔴 Missing: {', '.join(res['missing_ranked'][:10])}...")

        print(f"\n📋 SECTIONS")
        for sec, ok in res['sec_status'].items():
            print(f"  {'✅' if ok else '❌'} {sec}")

        print(f"\n✍️  LANGUAGE & IMPACT")
        print(f"  Verbs: {res['verb_count']} | Clichés: {res['cliche_count']} | Metrics: {res['metric_count']}")
        print(f"{W}\n")

    except Exception as e:
        print(f"FAILED: {e}")

if __name__ == "__main__":
    run_cv_audit(CV_PATH)
